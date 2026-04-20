from __future__ import annotations

import math
import statistics
import zipfile
import xml.etree.ElementTree as ET
from collections import Counter
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


WORKBOOK_PATH = Path("/Users/wawrzyk/Downloads/StudentsPerformance_excel_ukol (1) (1).xlsx")
OUTPUT_DIR = Path("/Users/wawrzyk/Desktop/Archive 2/report_output")
OUTPUT_DOCX = OUTPUT_DIR / "Semestralni_prace_StudentsPerformance.docx"
OFFICIAL_LOGO_PATH = OUTPUT_DIR / "osu_logo_official.png"
STUDENT_NAME = "Jan Wawrzyk"
SCHOOL_NAME = "Prirodovedecka fakulta, Ostravska univerzita"
FACULTY_COLOR = "00A67E"
NS = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}


def parse_sheet_rows() -> list[dict[str, str | int]]:
    def cell_value(cell: ET.Element) -> str | None:
        cell_type = cell.attrib.get("t")
        if cell_type == "inlineStr":
            return "".join(node.text or "" for node in cell.findall(".//m:t", NS))
        value = cell.find("m:v", NS)
        return None if value is None or value.text is None else value.text

    with zipfile.ZipFile(WORKBOOK_PATH) as archive:
        xml_root = ET.fromstring(archive.read("xl/worksheets/sheet1.xml"))
        rows = xml_root.find("m:sheetData", NS)
        assert rows is not None

        parsed_rows: list[dict[str, str | int]] = []
        for row in rows.findall("m:row", NS):
            values: dict[str, str | int] = {}
            for cell in row.findall("m:c", NS):
                ref = cell.attrib["r"]
                col = "".join(char for char in ref if char.isalpha())
                raw = cell_value(cell)
                values[col] = raw if raw is not None else ""
            parsed_rows.append(values)

    headers = [str(parsed_rows[0][col]) for col in ["A", "B", "C", "D", "E", "F", "G", "H"]]
    result: list[dict[str, str | int]] = []
    for row in parsed_rows[1:1001]:
        result.append(
            {
                headers[0]: str(row["A"]),
                headers[1]: str(row["B"]),
                headers[2]: str(row["C"]),
                headers[3]: str(row["D"]),
                headers[4]: str(row["E"]),
                headers[5]: int(str(row["F"])),
                headers[6]: int(str(row["G"])),
                headers[7]: int(str(row["H"])),
            }
        )
    return result


def mean(values: list[int]) -> float:
    return round(statistics.mean(values), 2)


def stdev(values: list[int]) -> float:
    return round(statistics.pstdev(values), 2)


def correlation(first: list[int], second: list[int]) -> float:
    mean_first = statistics.mean(first)
    mean_second = statistics.mean(second)
    covariance = sum((x - mean_first) * (y - mean_second) for x, y in zip(first, second)) / len(first)
    return round(covariance / (statistics.pstdev(first) * statistics.pstdev(second)), 3)


def prepare_stats(rows: list[dict[str, str | int]]) -> dict[str, object]:
    metrics = {
        "math": [int(row["math score"]) for row in rows],
        "reading": [int(row["reading score"]) for row in rows],
        "writing": [int(row["writing score"]) for row in rows],
    }

    gender_groups = {
        gender: [row for row in rows if row["gender"] == gender]
        for gender in ["female", "male"]
    }
    lunch_groups = {
        lunch: [row for row in rows if row["lunch"] == lunch]
        for lunch in ["standard", "free/reduced"]
    }
    prep_groups = {
        prep: [row for row in rows if row["test preparation course"] == prep]
        for prep in ["completed", "none"]
    }
    education_levels = [
        "master's degree",
        "bachelor's degree",
        "associate's degree",
        "some college",
        "high school",
        "some high school",
    ]
    education_groups = {
        level: [row for row in rows if row["parental level of education"] == level]
        for level in education_levels
    }
    race_counts = Counter(str(row["race_ethnicity"]) for row in rows)

    return {
        "count": len(rows),
        "metrics": metrics,
        "description": {
            "math": {"min": min(metrics["math"]), "max": max(metrics["math"]), "median": statistics.median(metrics["math"]), "mean": mean(metrics["math"]), "stdev": stdev(metrics["math"])},
            "reading": {"min": min(metrics["reading"]), "max": max(metrics["reading"]), "median": statistics.median(metrics["reading"]), "mean": mean(metrics["reading"]), "stdev": stdev(metrics["reading"])},
            "writing": {"min": min(metrics["writing"]), "max": max(metrics["writing"]), "median": statistics.median(metrics["writing"]), "mean": mean(metrics["writing"]), "stdev": stdev(metrics["writing"])},
        },
        "gender_summary": {
            gender: {
                "count": len(group),
                "math": mean([int(row["math score"]) for row in group]),
                "reading": mean([int(row["reading score"]) for row in group]),
                "writing": mean([int(row["writing score"]) for row in group]),
                "sum_writing": sum(int(row["writing score"]) for row in group),
            }
            for gender, group in gender_groups.items()
        },
        "lunch_summary": {
            lunch: {
                "count": len(group),
                "math": mean([int(row["math score"]) for row in group]),
                "reading": mean([int(row["reading score"]) for row in group]),
                "writing": mean([int(row["writing score"]) for row in group]),
            }
            for lunch, group in lunch_groups.items()
        },
        "prep_summary": {
            prep: {
                "count": len(group),
                "math": mean([int(row["math score"]) for row in group]),
                "reading": mean([int(row["reading score"]) for row in group]),
                "writing": mean([int(row["writing score"]) for row in group]),
            }
            for prep, group in prep_groups.items()
        },
        "education_summary": {
            level: {
                "count": len(group),
                "math": mean([int(row["math score"]) for row in group]),
                "reading": mean([int(row["reading score"]) for row in group]),
                "writing": mean([int(row["writing score"]) for row in group]),
            }
            for level, group in education_groups.items()
        },
        "race_counts": race_counts,
        "correlations": {
            "math_reading": correlation(metrics["math"], metrics["reading"]),
            "math_writing": correlation(metrics["math"], metrics["writing"]),
            "reading_writing": correlation(metrics["reading"], metrics["writing"]),
        },
    }


def configure_plot_style() -> None:
    plt.style.use("seaborn-v0_8-whitegrid")
    plt.rcParams["font.size"] = 10
    plt.rcParams["axes.titlesize"] = 13
    plt.rcParams["axes.labelsize"] = 10


def save_chart(path: Path, width: float = 8.5, height: float = 4.8) -> None:
    plt.tight_layout()
    plt.savefig(path, dpi=220, bbox_inches="tight")
    plt.close()


def create_logo_placeholder(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(4.2, 1.6))
    ax.axis("off")
    ax.add_patch(plt.Rectangle((0, 0), 1, 1, color="#003B6F", transform=ax.transAxes))
    ax.text(
        0.5,
        0.62,
        "LOGO FAKULTY",
        ha="center",
        va="center",
        fontsize=18,
        color="white",
        fontweight="bold",
        transform=ax.transAxes,
    )
    ax.text(
        0.5,
        0.28,
        "Nahraďte oficiálním logem školy",
        ha="center",
        va="center",
        fontsize=10,
        color="white",
        transform=ax.transAxes,
    )
    save_chart(path)


def create_infographic(path: Path, title: str, lines: list[str]) -> None:
    fig, ax = plt.subplots(figsize=(8.6, 4.8))
    ax.axis("off")
    ax.add_patch(plt.Rectangle((0, 0), 1, 1, color="#F4FBF8", transform=ax.transAxes))
    ax.add_patch(plt.Rectangle((0, 0.78), 1, 0.22, color=f"#{FACULTY_COLOR}", transform=ax.transAxes))
    ax.text(0.04, 0.89, title, ha="left", va="center", fontsize=18, color="white", fontweight="bold", transform=ax.transAxes)

    start_y = 0.67
    for index, line in enumerate(lines):
        y = start_y - (index * 0.12)
        ax.add_patch(plt.Circle((0.05, y), 0.012, color=f"#{FACULTY_COLOR}", transform=ax.transAxes))
        ax.text(0.08, y, line, ha="left", va="center", fontsize=12, color="#173A2F", transform=ax.transAxes)

    save_chart(path)


def create_charts(stats: dict[str, object]) -> dict[str, Path]:
    configure_plot_style()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    files = {
        "logo": OUTPUT_DIR / "logo_placeholder.png",
        "gender": OUTPUT_DIR / "graf_gender.png",
        "lunch": OUTPUT_DIR / "graf_lunch.png",
        "prep": OUTPUT_DIR / "graf_prep.png",
        "education": OUTPUT_DIR / "graf_education.png",
        "distribution": OUTPUT_DIR / "graf_distribution.png",
        "image_dataset": OUTPUT_DIR / "obrazek_dataset.png",
        "image_gender": OUTPUT_DIR / "obrazek_gender.png",
        "image_lunch": OUTPUT_DIR / "obrazek_lunch.png",
        "image_prep": OUTPUT_DIR / "obrazek_prep.png",
        "image_findings": OUTPUT_DIR / "obrazek_findings.png",
    }

    if OFFICIAL_LOGO_PATH.exists():
        files["logo"] = OFFICIAL_LOGO_PATH
    else:
        create_logo_placeholder(files["logo"])

    gender_summary = stats["gender_summary"]
    genders = ["female", "male"]
    x_positions = range(len(genders))
    bar_width = 0.24

    plt.figure(figsize=(8.5, 4.8))
    plt.bar([x - bar_width for x in x_positions], [gender_summary[g]["math"] for g in genders], width=bar_width, label="Matematika", color="#1f77b4")
    plt.bar(list(x_positions), [gender_summary[g]["reading"] for g in genders], width=bar_width, label="Čtení", color="#ff7f0e")
    plt.bar([x + bar_width for x in x_positions], [gender_summary[g]["writing"] for g in genders], width=bar_width, label="Psaní", color="#2ca02c")
    plt.xticks(list(x_positions), ["Ženy", "Muži"])
    plt.ylabel("Průměrné skóre")
    plt.title("Průměrné skóre podle pohlaví")
    plt.legend()
    save_chart(files["gender"])

    lunch_summary = stats["lunch_summary"]
    lunches = ["standard", "free/reduced"]
    plt.figure(figsize=(8.5, 4.8))
    plt.bar([x - bar_width for x in x_positions], [lunch_summary[g]["math"] for g in lunches], width=bar_width, label="Matematika", color="#1f77b4")
    plt.bar(list(x_positions), [lunch_summary[g]["reading"] for g in lunches], width=bar_width, label="Čtení", color="#ff7f0e")
    plt.bar([x + bar_width for x in x_positions], [lunch_summary[g]["writing"] for g in lunches], width=bar_width, label="Psaní", color="#2ca02c")
    plt.xticks(list(x_positions), ["Standard", "Free/reduced"])
    plt.ylabel("Průměrné skóre")
    plt.title("Průměrné skóre podle typu stravování")
    plt.legend()
    save_chart(files["lunch"])

    prep_summary = stats["prep_summary"]
    preps = ["completed", "none"]
    plt.figure(figsize=(8.5, 4.8))
    plt.bar([x - bar_width for x in x_positions], [prep_summary[g]["math"] for g in preps], width=bar_width, label="Matematika", color="#1f77b4")
    plt.bar(list(x_positions), [prep_summary[g]["reading"] for g in preps], width=bar_width, label="Čtení", color="#ff7f0e")
    plt.bar([x + bar_width for x in x_positions], [prep_summary[g]["writing"] for g in preps], width=bar_width, label="Psaní", color="#2ca02c")
    plt.xticks(list(x_positions), ["Kurz dokončen", "Bez kurzu"])
    plt.ylabel("Průměrné skóre")
    plt.title("Vliv přípravného kurzu na výsledky")
    plt.legend()
    save_chart(files["prep"])

    education_summary = stats["education_summary"]
    edu_labels = ["Master", "Bachelor", "Associate", "Some college", "High school", "Some high school"]
    edu_keys = [
        "master's degree",
        "bachelor's degree",
        "associate's degree",
        "some college",
        "high school",
        "some high school",
    ]
    plt.figure(figsize=(9.2, 5.0))
    plt.bar(edu_labels, [education_summary[key]["writing"] for key in edu_keys], color="#7f3c8d")
    plt.ylabel("Průměrné skóre v psaní")
    plt.title("Psaní podle vzdělání rodičů")
    plt.xticks(rotation=20, ha="right")
    save_chart(files["education"])

    description = stats["description"]
    plt.figure(figsize=(8.5, 4.8))
    subjects = ["Matematika", "Čtení", "Psaní"]
    means = [description["math"]["mean"], description["reading"]["mean"], description["writing"]["mean"]]
    errors = [description["math"]["stdev"], description["reading"]["stdev"], description["writing"]["stdev"]]
    plt.bar(subjects, means, yerr=errors, capsize=6, color=["#1f77b4", "#ff7f0e", "#2ca02c"])
    plt.ylabel("Skóre")
    plt.title("Průměr a variabilita výsledků")
    save_chart(files["distribution"])

    create_infographic(
        files["image_dataset"],
        "Profil datasetu",
        [
            f"Celkovy pocet studentu: {stats['count']}",
            f"Prumerny vysledek z matematiky: {description['math']['mean']} bodu",
            f"Prumerny vysledek ze cteni: {description['reading']['mean']} bodu",
            f"Prumerny vysledek z psani: {description['writing']['mean']} bodu",
        ],
    )
    create_infographic(
        files["image_gender"],
        "Srovnani podle pohlavi",
        [
            f"Zeny: cteni {gender_summary['female']['reading']} a psani {gender_summary['female']['writing']}",
            f"Muzi: matematika {gender_summary['male']['math']}",
            f"Rozdil mezi skupinami je nejvyraznejsi v jazykovych disciplinach",
            "Vysledky odpovidaji beznym trendum v mezinarodnich studiich",
        ],
    )
    create_infographic(
        files["image_lunch"],
        "Socioekonomicky kontext",
        [
            f"Standardni stravovani: matematika {lunch_summary['standard']['math']}",
            f"Free/reduced: matematika {lunch_summary['free/reduced']['math']}",
            f"Rozdil ve cteni: {round(lunch_summary['standard']['reading'] - lunch_summary['free/reduced']['reading'], 2)} bodu",
            "Typ stravovani zde funguje jako orientacni socialni ukazatel",
        ],
    )
    create_infographic(
        files["image_prep"],
        "Efekt pripravneho kurzu",
        [
            f"S kurzem: cteni {prep_summary['completed']['reading']}, psani {prep_summary['completed']['writing']}",
            f"Bez kurzu: cteni {prep_summary['none']['reading']}, psani {prep_summary['none']['writing']}",
            "Priprava zlepsuje vysledky ve vsech trech castich testu",
            "Nejvetsi rozdil je patrny v psani",
        ],
    )
    create_infographic(
        files["image_findings"],
        "Klicova zjisteni analyzy",
        [
            f"Nejsilnejsi korelace: cteni a psani ({stats['correlations']['reading_writing']})",
            "Vyssi vzdelani rodicu souvisi s lepsimi vysledky studentu",
            "Zeny prekonavaji muze ve cteni a psani",
            "Muzi maji v souboru vyssi matematicky prumer",
        ],
    )

    return files


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, fld_end])


def clear_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ""
    section.even_page_footer.paragraphs[0].text = ""


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Obsah se načte po aktualizaci polí ve Wordu."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def create_heading_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_num_id = 90
    num_id = 90

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level_type)

    for level, level_text in enumerate(["%1", "%1.%2", "%1.%2.%3"]):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)

        lvl_text_element = OxmlElement("w:lvlText")
        lvl_text_element.set(qn("w:val"), level_text)
        lvl.append(lvl_text_element)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(360 * (level + 1)))
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)

        abstract_num.append(lvl)

    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_ref = OxmlElement("w:abstractNumId")
    abstract_num_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_num_ref)
    numbering.append(num)

    return num_id


def apply_heading_numbering(paragraph, numbering_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing_num_pr = p_pr.find(qn("w:numPr"))
    if existing_num_pr is not None:
        p_pr.remove(existing_num_pr)

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level - 1))
    num_pr.append(ilvl)

    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(numbering_id))
    num_pr.append(num_id)

    p_pr.append(num_pr)


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(document: Document, title: str, level: int, numbering_id: int | None = None) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    if numbering_id is not None:
        apply_heading_numbering(paragraph, numbering_id, level)
    run = paragraph.add_run(title)
    if level == 1:
        run.font.all_caps = True


def apply_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1.font.size = Pt(15)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor.from_string(FACULTY_COLOR)

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)

    heading3 = document.styles["Heading 3"]
    heading3.font.name = "Calibri"
    heading3.font.size = Pt(11)
    heading3.font.bold = False
    heading3.font.color.rgb = RGBColor(0, 0, 0)

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(10)
    caption.font.italic = True


def add_paragraph(document: Document, text: str, italic: bool = False, center: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = italic


def add_caption(paragraph, label: str, text: str) -> None:
    paragraph.add_run(f"{label} ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {label} \\* ARABIC "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, placeholder, fld_end])
    paragraph.add_run(f": {text}")


def add_generated_list(document: Document, label: str, empty_text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\h \\z \\c "{label}" '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = empty_text
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, placeholder, fld_end])


def add_figure(document: Document, image_path: Path, caption_text: str, width_inches: float = 6.2, label: str = "Obrazek") -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))

    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(caption, label, caption_text)


def add_table_caption(document: Document, text: str, label: str = "Tabulka") -> None:
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(caption, label, text)


def build_report(stats: dict[str, object], charts: dict[str, Path]) -> None:
    document = Document()
    document.settings.odd_and_even_pages_header_footer = True
    set_update_fields_on_open(document)

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True

    apply_document_styles(document)
    heading_numbering_id = create_heading_numbering(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run().add_picture(str(charts["logo"]), width=Inches(3.2))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ANALYZA STUDIJNICH VYSLEDKU STUDENTU")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(FACULTY_COLOR)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Semestralni prace do textoveho procesoru")
    subtitle_run.font.size = Pt(13)

    for line in [
        "",
        f"Student: {STUDENT_NAME}",
        f"Skola: {SCHOOL_NAME}",
        "Akademicky rok: 2025/2026",
        "",
    ]:
        add_paragraph(document, line, center=True)

    document.add_page_break()

    add_heading(document, "Obsah", 1)
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)

    main_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    main_section.different_first_page_header_footer = True
    main_section.footer.is_linked_to_previous = False
    main_section.even_page_footer.is_linked_to_previous = False
    set_page_number_start(main_section, 1)

    odd_footer = main_section.footer.paragraphs[0]
    odd_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_page_number(odd_footer)

    even_footer = main_section.even_page_footer.paragraphs[0]
    even_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(even_footer)

    description = stats["description"]
    gender_summary = stats["gender_summary"]
    lunch_summary = stats["lunch_summary"]
    prep_summary = stats["prep_summary"]
    education_summary = stats["education_summary"]
    race_counts = stats["race_counts"]
    correlations = stats["correlations"]

    add_heading(document, "Uvod do analyzy vysledku studentu", 1, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        "Tato prace se zabyva analyzou datoveho souboru Students Performance in Exams. Cilem je prehledne ukazat, jakych vysledku studenti dosahuji v matematice, cteni a psani, a soucasne vyhodnotit, jak se do jejich uspesnosti promita pohlavi, rodinne zazemi, typ stravovani nebo absolvovani pripravneho kurzu. Pred samotnou interpretaci vysledku budou nejprve predstavena pouzita data a metodika zpracovani, nasledne budou popsany hlavni statisticke ukazatele a porovnany rozdily mezi jednotlivymi skupinami studentu.",
    )
    add_paragraph(
        document,
        "V soucasnem vyzkumu vzdelavani se opakovane ukazuje, ze vykon studentu neni dan pouze jejich individualnimi schopnostmi, ale take rodinnym zazemim, socioekonomickym statusem a podpurnymi mechanismy skoly i domova. OECD dlouhodobe upozornuje, ze spravedlnost ve vzdelavani a pristup ke studijnim zdrojum vyznamne ovlivnuji vysledky zaku [C2][C3]. Podobne zpravy UNESCO zduraznuji, ze rozdily mezi skupinami studentu se projevuji jak v matematice, tak ve ctenarske gramotnosti, pricemz cast rozdilu souvisi s rodinnym a socialnim kontextem [C4][C5].",
    )
    add_figure(document, charts["image_dataset"], "Souhrn nejdulezitejsich charakteristik datasetu studentu.")

    add_heading(document, "Metodika a zdroj dat", 2, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        "Datovy soubor vychazi z bezne dostupneho datasetu Students Performance in Exams, ktery je casto vyuzivan pro vyuku zakladni datove analyzy a statistiky [C1]. V Excelu byly pripraveny souhrnne funkce, kontingencni tabulka a alespon jeden graf. Pro potreby teto prace byla navic provedena interpretace klicovych ukazatelu, zejmena minima, maxima, medianu a prumeru, a dale porovnani vybranych skupin studentu podle pohlavi, typu stravovani, pripravy na test a vzdelani rodicu.",
    )
    add_heading(document, "Charakteristika promennych", 3, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        f"Zakladni analyza ukazala, ze prumerne skore v matematice dosahuje {description['math']['mean']}, ve cteni {description['reading']['mean']} a v psani {description['writing']['mean']} bodu. Median je u matematiky {description['math']['median']}, u cteni {description['reading']['median']} a u psani {description['writing']['median']}. Z hlediska variability se nejvetsi smerodatna odchylka vyskytuje u psani ({description['writing']['stdev']}), zatimco cteni je o neco stabilnejsi ({description['reading']['stdev']}).",
    )

    add_heading(document, "Struktura sledovaneho souboru", 2, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        f"Soubor obsahuje celkem {stats['count']} studentu. Nejvice zaznamu spada do etnicke skupiny C ({race_counts['group C']} studentu), nasledovane skupinou D ({race_counts['group D']}) a skupinou B ({race_counts['group B']}). Zastoupeni zen je mirne vyssi ({gender_summary['female']['count']}) nez zastoupeni muzu ({gender_summary['male']['count']}). Tento fakt je dulezity zejmena pro interpretaci souhrnnych prumeru a vizualizaci podle pohlavi.",
    )
    add_figure(document, charts["image_gender"], "Prehled rozdilu mezi studentkami a studenty v jednotlivych typech skore.")

    add_heading(document, "Popisna statistika vysledku", 1, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        "Popisna statistika potvrzuje, ze nejvyssi dosazene skore ve vsech trech testovych oblastech bylo 100 bodu, zatimco nejmensi hodnoty se pohybuji od 0 do 17 bodu. To znamena, ze dataset zachycuje velmi siroke spektrum vykonnosti od slabych po excelentni vysledky. Zajimave je, ze prumerne hodnoty cteni a psani jsou vyssi nez u matematiky, coz odpovida castym zjistenim mezinarodnich srovnavacich studii [C2][C4].",
    )
    add_paragraph(
        document,
        f"Z hlediska vzajemnych vztahu mezi predmety vykazuje nejsilnejsi korelaci dvojice cteni a psani ({correlations['reading_writing']}). O neco nizsi, ale stale velmi vyrazna je korelace mezi matematikou a ctenim ({correlations['math_reading']}) a mezi matematikou a psanim ({correlations['math_writing']}). Lze tedy predpokladat, ze studenti s lepsimi vysledky v jedne oblasti maji casto dobre vysledky i v ostatnich oblastech, i kdyz matematika se od jazykove orientovanych disciplin casto mirne odlisuje [C6].",
    )
    add_figure(document, charts["distribution"], "Prumer a variabilita vysledku v matematice, cteni a psani.", label="Graf")
    add_paragraph(
        document,
        "Graf 1 shrnuje nejen prumerne hodnoty, ale neprimo i rozptyl vysledku. Zejmena u matematiky je patrne, ze cast studentu dosahuje vyrazne podprumernych vysledku, coz muze signalizovat rozdily v predchozi priprave nebo v rodinnych podminkach pro studium.",
    )

    add_heading(document, "Rozdily podle pohlavi", 2, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        f"Pri porovnani podle pohlavi se ukazuje zajimavy vzorec. Zeny dosahuji lepsich vysledku ve cteni ({gender_summary['female']['reading']}) a psani ({gender_summary['female']['writing']}), zatimco muzi vykazuji vyssi prumer v matematice ({gender_summary['male']['math']}). Tento rozdil je dobre znamym jevem i v mezinarodnich datech a byva vysvetlovan kombinaci socializace, vyukovych strategii a odlisnych preferenci pri reseni ukolu [C2][C4].",
    )
    add_heading(document, "Interpretace genderovych rozdilu", 3, numbering_id=heading_numbering_id)
    add_figure(document, charts["gender"], "Prumerne skore podle pohlavi.", label="Graf")
    add_paragraph(
        document,
        f"Graf 2 navazuje na vysledky z predchozi excelove analyzy a potvrzuje, ze zeny v datasetu vyrazne prevysuji muze v jazykove orientovanych oblastech. Naopak matematicke skore je vyssi u muzu ({gender_summary['male']['math']}) nez u zen ({gender_summary['female']['math']}). Tyto zavery je vsak vhodne vnimat jako deskriptivni, nikoliv jako dukaz pricinnosti.",
    )

    add_heading(document, "Rozdily podle socioekonomickych ukazatelu", 2, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        f"Jako orientacni ukazatel socioekonomicke situace lze v datasetu vyuzit promennou typ stravovani. Studenti se standardnim stravovanim dosahuji ve vsech trech castech testu vyssich prumeru nez studenti se stravovanim free/reduced. V matematice jde o rozdil {round(lunch_summary['standard']['math'] - lunch_summary['free/reduced']['math'], 2)} bodu, ve cteni o {round(lunch_summary['standard']['reading'] - lunch_summary['free/reduced']['reading'], 2)} bodu a v psani o {round(lunch_summary['standard']['writing'] - lunch_summary['free/reduced']['writing'], 2)} bodu. Tento vysledek je v souladu s literaturem o vlivu socioekonomickeho zazemi na skolni vykon [C3][C5].",
    )
    add_figure(document, charts["lunch"], "Prumerne skore podle typu stravovani.", label="Graf")
    add_paragraph(
        document,
        "Rozdily v Grafu 3 jsou natolik vyrazne, ze je lze povazovat za jeden z nejdulezitejsich vzorcu v celem souboru. Nejde pouze o matematiku, ale o konzistentni nevyhodu ve vsech merenech oblastech. Tento obraz podporuje nazor, ze socialni podminky a dostupnost podpory mimo skolu maji na studijni uspech realny dopad.",
    )
    add_figure(document, charts["image_lunch"], "Infograficky souhrn rozdilu podle typu stravovani.")

    add_heading(document, "Vliv pripravy a rodinneho prostredi", 1, numbering_id=heading_numbering_id)
    add_heading(document, "Pripravny kurz", 2, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        f"Absolvovani pripravneho kurzu je spojeno s lepsimi vysledky ve vsech trech testech. Studenti, kteri kurz dokoncili, maji prumer v matematice {prep_summary['completed']['math']}, ve cteni {prep_summary['completed']['reading']} a v psani {prep_summary['completed']['writing']}. Ve srovnani se studenty bez kurzu je jejich vykon vyssi o {round(prep_summary['completed']['math'] - prep_summary['none']['math'], 2)} bodu v matematice, o {round(prep_summary['completed']['reading'] - prep_summary['none']['reading'], 2)} bodu ve cteni a o {round(prep_summary['completed']['writing'] - prep_summary['none']['writing'], 2)} bodu v psani.",
    )
    add_figure(document, charts["prep"], "Vliv pripravneho kurzu na prumerne vysledky.", label="Graf")
    add_paragraph(
        document,
        "Graf 4 ilustruje, ze pozitivni efekt pripravy je patrny zvlast ve ctenarske a pisemne casti. Odborna literatura upozornuje, ze organizovana priprava, pravidelne opakovani a zpetna vazba mohou zvysovat nejen znalosti, ale i jistotu pri testovani [C6][C7].",
    )
    add_figure(document, charts["image_prep"], "Shrnuti rozdilu mezi studenty s pripravnym kurzem a bez nej.")

    add_heading(document, "Vzdelani rodicu", 2, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        "Dalsim sledovanym faktorem je nejvyssi dosazene vzdelani rodicu. Obecne plati, ze studenti, jejichz rodice maji vyssi vzdelani, vykazuji lepsi vysledky. Nejvyssi prumer v psani byl zaznamenan u skupiny master's degree, nasledovane bachelor's degree. Naopak nejnizsi hodnoty se objevuji u skupin high school a some high school. Tento vzorec odpovida poznatkum OECD o mezigeneracnim prenosu vzdelavacich sanci a roli rodinneho kulturniho kapitalu [C3][C8].",
    )
    add_heading(document, "Rodinne zazemi a skore studentu", 3, numbering_id=heading_numbering_id)
    add_figure(document, charts["education"], "Prumerne skore v psani podle vzdelani rodicu.", label="Graf")
    masters_writing = education_summary["master's degree"]["writing"]
    some_high_school_writing = education_summary["some high school"]["writing"]
    add_paragraph(
        document,
        f"Napriklad skupina s rodici s magisterskym vzdelanim dosahuje v psani prumerne hodnoty {masters_writing}, zatimco skupina some high school pouze {some_high_school_writing}. Rozdil mezi obema skupinami tedy presahuje deset bodu, coz nelze povazovat za zanedbatelne.",
    )
    add_figure(document, charts["image_findings"], "Klicova zjisteni analyzy v podobe souhrnneho obrazku.")

    add_heading(document, "Kontingencni tabulka a interpretace", 1, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        "V predchozim excelovem ukolu byla pripravena souhrnna kontingencni tabulka seskupena podle pohlavi. Tabulka zachycuje soucet bodu ze psani, prumerne skore z matematiky a pocet studentu. Jeji vyhodou je rychly prehled o rozdilech mezi obema skupinami a soucasne i kontrola velikosti vzorku.",
    )

    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["Pohlavi", "Suma Writing (body)", "Prumer Math (body)", "Pocet studentu"]
    values = [
        ["female", str(gender_summary["female"]["sum_writing"]), str(gender_summary["female"]["math"]), str(gender_summary["female"]["count"])],
        ["male", str(gender_summary["male"]["sum_writing"]), str(gender_summary["male"]["math"]), str(gender_summary["male"]["count"])],
        ["Celkem", str(gender_summary["female"]["sum_writing"] + gender_summary["male"]["sum_writing"]), str(description["math"]["mean"]), str(stats["count"])],
    ]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_index, row_values in enumerate(values, start=1):
        for col_index, cell_value in enumerate(row_values):
            table.rows[row_index].cells[col_index].text = cell_value
    add_table_caption(document, "Kontingencni tabulka podle pohlavi odpovidajici predchozi excelove analyze.")

    add_paragraph(
        document,
        f"Z Tabulky 1 vyplyva, ze zeny tvori mirne pocetnejsi skupinu ({gender_summary['female']['count']} studentu), ale soucasne dosahuji vyssiho souctu bodu v psani ({gender_summary['female']['sum_writing']}) nez muzi ({gender_summary['male']['sum_writing']}). V matematice je naopak prumerne uspesnejsi muzska skupina ({gender_summary['male']['math']} bodu oproti {gender_summary['female']['math']} bodu).",
    )
    add_paragraph(
        document,
        "Kontingencni tabulka tak dobre potvrzuje, ze rozdily mezi skupinami nejsou jednorozmerne. Podle zvoleneho ukazatele se meni i interpretace: zatimco celkovy objem bodu v psani nahrava zenam, matematicky prumer vychazi lepe muzum. Prave proto je vhodne kombinovat kontingencni tabulky s grafy a slovni interpretaci.",
    )

    add_heading(document, "Diskuse a zaver", 1, numbering_id=heading_numbering_id)
    add_paragraph(
        document,
        "Analyza ukazala, ze dataset Students Performance in Exams obsahuje nekolik stabilnich vzorcu. Nejvyraznejsi je silna souvislost mezi ctenim a psanim, dale rozdily podle pohlavi a velmi vyznamne odchylky podle typu stravovani a absolvovani pripravneho kurzu. Vysledky tedy podporuji tezi, ze skolni uspech je ovlivnen jak individualnimi studijnimi schopnostmi, tak sirsim socialnim a rodinnym kontextem [C2][C3][C8].",
    )
    add_paragraph(
        document,
        "Z praktickeho hlediska lze data interpretovat tak, ze podpora studentu nema smerovat pouze na vykon samotny, ale i na podminky, ktere vykon ovlivnuji. Vhodna intervence muze zahrnovat dostupne priprave kurzy, podporu ctenarske gramotnosti a vcasnou pomoc studentum ze socialne znevyhodneneho prostredi. Takovy pristup odpovida doporucenim OECD i UNESCO pro zmirnovani rozdilu ve vzdelavacich vysledcich [C4][C5].",
    )
    add_paragraph(
        document,
        "Soucasne je nutne zduraznit, ze jde o deskriptivni analyzu jedne konkretni datove sady. Bez dalsich promennych a bez inferencnich metod nelze tvrdit, ze nektery z pozorovanych vztahu je primo pricinny. I presto je analyza dostatecne prukazna pro ukazku prace s tabulkami, grafy a slovni interpretaci vysledku v textovem procesoru.",
    )

    lists_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    lists_section.different_first_page_header_footer = True
    clear_footer(lists_section)

    add_heading(document, "Seznam obrazku", 1)
    add_generated_list(document, "Obrazek", "Seznam obrazku se doplni po otevreni dokumentu ve Wordu.")
    add_heading(document, "Seznam grafu", 1)
    add_generated_list(document, "Graf", "Seznam grafu se doplni po otevreni dokumentu ve Wordu.")
    add_heading(document, "Seznam tabulek", 1)
    add_generated_list(document, "Tabulka", "Seznam tabulek se doplni po otevreni dokumentu ve Wordu.")

    bibliography_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    bibliography_section.different_first_page_header_footer = True
    clear_footer(bibliography_section)

    add_heading(document, "Pouzita literatura a zdroje", 1)
    add_paragraph(document, "[C1] spscientist. Students Performance in Exams. Kaggle. Dostupne z: https://www.kaggle.com/datasets/spscientist/students-performance-in-exams/data")
    add_paragraph(document, "[C2] OECD. PISA 2022 Results (Volume I): The State of Learning and Equity in Education. Paris: OECD Publishing, 2023.")
    add_paragraph(document, "[C3] OECD. Equity in Education: Breaking Down Barriers to Social Mobility. Paris: OECD Publishing, 2018.")
    add_paragraph(document, "[C4] OECD. PISA 2022 Results (Volume II): Learning During - and From - Disruption. Paris: OECD Publishing, 2023.")
    add_paragraph(document, "[C5] UNESCO. Global education monitoring report, 2024/5, Leadership in education: lead for learning. Paris: UNESCO, 2024.")
    add_paragraph(document, "[C6] Adesope, O. O.; Trevisan, D. A.; Sundararajan, N. Rethinking the Use of Tests: A Meta-Analysis of Practice Testing. Review of Educational Research, 2017.")
    add_paragraph(document, "[C7] Hao, J. et al. The Impact of Test Preparation on Performance of Large-Scale Educational Tests: A Meta-analysis of Experimental Studies. 2025.")
    add_paragraph(document, "[C8] OECD. Let's Read Them a Story! The Parent Factor in Education. Paris: OECD Publishing, 2012.")
    add_paragraph(
        document,
        "Poznamka: Pozadavane citace je potreba ve Wordu vlozit pres CitacePRO. V tomto souboru jsou pripraveny textove zastupce [C1]-[C8], ktere lze jednoduse nahradit formalnimi citacemi a automaticky generovanym seznamem literatury.",
        italic=True,
    )

    document.save(OUTPUT_DOCX)


def main() -> None:
    rows = parse_sheet_rows()
    stats = prepare_stats(rows)
    charts = create_charts(stats)
    build_report(stats, charts)
    print(f"Created: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
